import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml import OxmlElement
import copy

import os
import re
from datetime import datetime
from pathlib import Path

import shutil

# import datetime


from document_format_setter import DocumentFormatSetter


#从docx文档中读取全部文本
def get_text(file_name):
    '''从docx文档中读取全部文本'''
    doc = Document(file_name)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def docx_old_to_docx_new(input_data):
    '''
    # 导入需要修改的文件的文本内容
    # 生成一个新的 docx 文档，防止被原文档的设置干扰
    '''
    if isinstance(input_data, str):
        if os.path.isfile(input_data) and input_data.endswith('.docx'):
            doc_in_text = get_text(input_data)
        else:
            doc_in_text = input_data
    elif isinstance(input_data, docx.document.Document):
        doc_in_text = '\n'.join([p.text for p in input_data.paragraphs])
    else:
        raise ValueError("输入参数类型不正确，应为 str（docx 文件路径）或 docx.document.Document 类型")

    document = Document()
    # 将导入的文件文本，分成各行，分别生成不同的段落。不分的话，是在同一个段落中。
    list_text = doc_in_text.split('\n')  # 不小心实现了删除全部空行
    for text in list_text:
        if text.strip():  # 增加判断，去除空行
            document.add_paragraph(text)

    return document


# 将txt文件转换为docx文件
def txt_to_docx(txt_file, docx_file):
    # 打开txt文件
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 将txt内容按照换行符进行分割,确保换行符不变。
    lines = content.split('\n')

    # 创建一个新的Word文档
    doc = Document()

    # 将每一行文本添加到docx文档中
    for line in lines:
        doc.add_paragraph(line)

    # 保存为docx文件
    doc.save(docx_file)
    return doc


# 将指定文件夹中的txt文件转换为docx文件
def txt_to_docx_all(folder_path_in):
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.txt'):
            txt_path_in = os.path.join(folder_path_in, file_name_in)
            # docxs.append(docx_path_in)
            docx_path_out = txt_path_in[:-3] + 'docx'

            txt_to_docx(txt_path_in, docx_path_out)



# 将txt文件转换为docx文件
def txt_to_docx_converted(txt_file):
    """
    打开txt文件，将内容转换为docx格式

    参数：
    txt_file (str): txt 文件路径

    返回：
    docx.Document: 转换后的 Word 文档对象
    """
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 将txt内容按照换行符进行分割, 确保换行符不变。
    lines = content.split('\n')

    # 创建一个新的 Word 文档
    doc = Document()

    # 将每一行文本添加到 docx 文档中
    for line in lines:
        doc.add_paragraph(line)

    return doc

def str_to_docx(input_str):
    """
    将输入字符串转换为docx文档
    有多行的，要分行添加到docx文档的段落中

    参数：
    input_str (str): 输入字符串

    返回：
    docx.Document: 转换后的 Word 文档对象
    """
    doc = Document()
    paragraphs = input_str.split('\n')
    '''当使用 input_str.split('\n') 来分割字符串时，
    如果存在连续的两个 '\n\n' ，
    那么会将这两个连续的换行符视为一个分隔符，
    分割出来的段落列表中，
    在对应位置会是一个空字符串元素。
    '''
    for paragraph in paragraphs:
        if paragraph:
            doc.add_paragraph(paragraph)
    return doc


def generate_new_file_path(original_path, new_content='', new_suffix=None, add_timestamp=False):
    """
    根据原始文件路径生成一个新的文件路径，允许修改后缀。

    :param original_path: 原始文件的绝对路径。
    :param new_suffix: 新文件的后缀，包括点（例如 '.txt' 而不是 'txt'）。
    :param add_timestamp: 是否在文件名后添加当前时间戳，默认为False。
    :return: 新文件的绝对路径。
    """
    # 创建Path对象
    original_file = Path(original_path)

    # 如果原始路径不是文件，抛出异常
    if not original_file.is_file():
        raise ValueError(f"The provided path '{original_path}' is not a file.")

    # 获取文件的目录和文件名
    directory = original_file.parent
    filename_without_extension = original_file.stem
    file_extension = original_file.suffix

    # 如果没有提供新后缀，使用原始文件的后缀
    if new_suffix is None:
        new_suffix = file_extension

    # 构建新文件名
    if add_timestamp:
        # 添加时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_filename = f"{filename_without_extension}_{new_content}_{timestamp}{new_suffix}"
    else:
        # 不添加时间戳
        new_filename = f"{filename_without_extension}_{new_content}{new_suffix}"

    # 创建新文件的完整路径
    new_file_path = directory / new_filename

    return str(new_file_path)

def file_to_docx(input_file, docx_path_out, hides_first=[0, 1],have_inscribe = True,date_and_name_under_title_flag = False,date_and_name_under_title_numbers = [1]):
    """
    根据输入文件的类型进行处理，如果是docx文件则直接处理，
    如果是txt文件则转换为docx后处理，如果是字符串则直接处理

    参数：
    input_file (str): 输入的文件路径或字符串
    docx_path_out (str): 输出的docx文件路径
    hides_first (list, 可选): 一些控制参数，默认为 [0, 1]
    默认有主送、落款
    返回：
    None
    """
    if input_file.endswith('.txt'):
        # 将txt文件转换为docx格式并处理
        doc = txt_to_docx_converted(input_file)
    elif input_file.endswith('.docx'):
        # 处理输入为docx文件的情况
        doc = Document(input_file)
    elif isinstance(input_file, str):
        # 处理输入为字符串的情况
        doc = str_to_docx(input_file)
    else:
        # 处理输入为docx文件的情况
        doc = Document(input_file)

    doc = docx_old_to_docx_new(doc)
    # # 设置页面
    # pages_set(doc)
    # # 设置页边距，可以从参数文件中读取，不用修改函数，尝试使用 JSON
    # sections_set(doc, left=2.8, right=2.6, top=3.7, bottom=3.5)
    # # 设置段落边距等
    # para_set_indent(doc)
    #
    # # 设置正文，全部先设置为正文格式，再逐个调整各级标题
    # paragraphs_set_all(doc)
    # # 设置各级标题
    # paragraphs_set_hides_first(doc, lines=hides_first)
    # paragraphs_set_hides_second(doc)
    # paragraphs_set_hides_third(doc)
    #
    # #如果有标题下面的日期和姓名，则改之
    # if date_and_name_under_title_flag:
    #     paragraphs_set_date_and_name_under_title(doc,lines = date_and_name_under_title_numbers)
    #
    # # 设置主送机关和落款单位、时间,默认为真，执行之
    # if have_inscribe: # = True
    #     paragraphs_set_inscribe(doc)
    #
    # # 添加页码，设置，如果页面超过 2 页，则添加，另外，如果
    # InsertPageNumber(doc)

    document_current = DocumentFormatSetter(doc)
    doc_new = document_current.document_to_offical_format()
    # 保存为新的文件
    doc_new.save(docx_path_out)
    return doc_new


file_path_in = r'D:\001git_yuxingruyu\all_to_docx_office_format\test_folder\test_input.txt'
file_path_out =  generate_new_file_path(file_path_in, new_content='普通公文格式',new_suffix='.docx', add_timestamp=True)
                        # print(f'{path_in}\n{path_out}\n\n')
                        # have_inscribe = False 代表尚有主送和落款
file_to_docx(file_path_in, file_path_out, hides_first=[0, 1], have_inscribe=False,
                                     date_and_name_under_title_flag=False, date_and_name_under_title_numbers=[1])