# Create:20241226
#modify codes using class and create testing code

from docx import Document

from docx.shared import Cm, Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml import OxmlElement
import os
import re
from pathlib import Path
import datetime
from docx import Document, document

class DocumentConverter:

    def __init__(self):
        self.document = None

    def docx_old_to_docx_new(self, input_data):
        """
        将旧的docx文件内容导入，生成一个新的docx文档，防止被原文档的设置干扰。
        也能处理直接传入字符串内容的情况（当作文档内容来处理）。

        参数:
        input_data (str 或 docx.document.Document): 可以是docx文件路径字符串，也可以是已经打开的Document对象，或者是纯文本内容字符串

        返回:
        docx.document.Document: 生成的新的docx文档对象

        异常处理:
        - 如果输入的文件路径不存在，抛出FileNotFoundError异常。
        - 如果打开文件时出现其他I/O异常，向上传递相应的I/O异常。
        """
        if isinstance(input_data, str):
            if not os.path.exists(input_data):
                raise FileNotFoundError(f"指定的文件 {input_data} 不存在")
            try:
                if input_data.endswith('.docx'):
                    doc_in_text = self.get_text(input_data)
                else:
                    doc_in_text = input_data
            except IOError as e:
                raise IOError(f"读取文件 {input_data} 时出现错误: {e}")
        elif isinstance(input_data, document.Document):
            doc_in_text = '\n'.join([p.text for p in input_data.paragraphs])
        else:
            raise ValueError("输入参数类型不正确，应为 str（docx 文件路径）或 docx.document.Document 类型")

        self.document = Document()
        list_text = doc_in_text.split('\n')
        for text in list_text:
            if text.strip():
                self.document.add_paragraph(text)

        return self.document

    def txt_to_docx(self, txt_file, docx_file):
        """
        将给定的txt文件内容转换为docx文件格式并保存。

        参数:
        txt_file (str): 要转换的txt文件路径
        docx_file (str): 转换后保存的docx文件路径

        返回:
        docx.document.Document: 转换后的docx文档对象

        异常处理:
        - 如果输入的txt文件路径不存在，抛出FileNotFoundError异常。
        - 如果保存docx文件时出现权限问题或其他I/O异常，向上传递相应的异常。
        """
        if not os.path.exists(txt_file):
            raise FileNotFoundError(f"指定的txt文件 {txt_file} 不存在")
        try:
            with open(txt_file, 'r', encoding='utf-8') as f:
                content = f.read()
            lines = content.split('\n')
            doc = Document()
            for line in lines:
                doc.add_paragraph(line)
            doc.save(docx_file)
            return doc
        except IOError as e:
            raise IOError(f"在转换或保存文件时出现错误: {e}")

    def txt_to_docx_all(self, folder_path_in):
        """
        将指定文件夹中的所有txt文件转换为docx文件。

        参数:
        folder_path_in (str): 包含txt文件的文件夹路径

        异常处理:
        - 如果输入的文件夹路径不存在，抛出FileNotFoundError异常。
        - 如果遍历文件夹下文件时出现权限问题等导致无法访问文件，向上传递相应的异常。
        """
        if not os.path.exists(folder_path_in):
            raise FileNotFoundError(f"指定的文件夹 {folder_path_in} 不存在")
        try:
            for file_name_in in os.listdir(folder_path_in):
                if file_name_in.endswith('.txt'):
                    txt_path_in = os.path.join(folder_path_in, file_name_in)
                    docx_path_out = txt_path_in[:-3] + 'docx'
                    self.txt_to_docx(txt_path_in, docx_path_out)
        except OSError as e:
            raise OSError(f"遍历文件夹 {folder_path_in} 时出现错误: {e}")

    def get_text(self, file_name):
        """
        从docx文档中读取全部文本内容。

        参数:
        file_name (str): docx文档的文件路径

        返回:
        str: 文档中的全部文本内容，以换行符连接各段落文本

        异常处理:
        - 如果文件不存在，抛出FileNotFoundError异常。
        - 如果读取文件过程中出现其他I/O异常，向上传递相应的I/O异常。
        """
        if not os.path.exists(file_name):
            raise FileNotFoundError(f"指定的docx文件 {file_name} 不存在")
        try:
            doc = Document(file_name)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except IOError as e:
            raise IOError(f"读取docx文件 {file_name} 时出现错误: {e}")


