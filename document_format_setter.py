# Create:20241226
# modify codes using class and create testing code

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml import OxmlElement
import copy

import os
import re

class DocumentFormatSetter:
    def __init__(self, document):
        self.document = document

    def pages_set(self, page_width=21, page_height=29.7, orientation=WD_ORIENT.PORTRAIT):
        """
        设置文档的纸张大小和页面方向。

        参数:
        page_width (float, 可选): 纸张宽度，单位为厘米，默认值为21
        page_height (float, 可选): 纸张高度，单位为厘米，默认值为29.7
        orientation (WD_ORIENT, 可选): 页面方向，默认值为纵向（WD_ORIENT.PORTRAIT）

        返回:
        None

        异常处理:
        - 如果设置页面属性时出现底层库相关的异常，向上传递异常，方便排查问题。
        """
        try:
            section = self.document.sections[-1]  # 创建一个新的section对象
            section.page_width = Cm(page_width)# 设置纸张大小为A4
            section.page_height = Cm(page_height)# 设置纸张大小为A4
            section.orientation = orientation# 设置页面方向为纵向
        except Exception as e:
            raise Exception(f"设置页面属性时出现错误: {e}")

    def sections_set(self, left=2.8, right=2.6, top=3.7, bottom=3.5):
        """
        设置文档的页边距。

        参数:
        left (float, 可选): 左边距，单位为厘米，默认值为2.8
        right (float, 可选): 右边距，单位为厘米，默认值为2.6
        top (float, 可选): 上边距，单位为厘米，默认值为3.7
        bottom (float, 可选): 下边距，单位为厘米，默认值为3.5

        返回:
        None

        异常处理:
        - 若在设置页边距过程中出现异常（如底层库不支持的操作等），向上抛出异常。
        """
        try:
            sections = self.document.sections
            for section in sections:
                section.left_margin = Cm(left)
                section.right_margin = Cm(right)
                section.top_margin = Cm(top)
                section.bottom_margin = Cm(bottom)
        except Exception as e:
            raise Exception(f"设置页边距时出现错误: {e}")

    def para_set_indent(
            self,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            left_indent=0,
            right_indent=0,
            space_before=0,
            space_after=0,
            line_spacing=28,
            first_line_indent=28
    ):
        """
        设置文档段落的缩进、间距和行间距等格式。

        参数:
        alignment (WD_ALIGN_PARAGRAPH, 可选): 段落对齐方式，默认值为两端对齐（WD_ALIGN_PARAGRAPH.JUSTIFY）
        left_indent (float, 可选): 左缩进，单位为磅，默认值为0
        right_indent (float, 可选): 右缩进，单位为磅，默认值为0
        space_before (float, 可选): 段前间距，单位为磅，默认值为0
        space_after (float, 可选): 段后间距，单位为磅，默认值为0
        line_spacing (float, 可选): 行间距，单位为磅，默认值为28
        first_line_indent (float, 可选): 首行缩进，单位为磅，默认值为28

        返回:
        None

        异常处理:
        - 若设置段落格式时出现问题（例如非法的参数值、底层库错误等），向上抛出异常以便处理。
        设置段落间距，默认使用公文的
        """
        try:
            for para in self.document.paragraphs:   #遍历 document 中的所有段落
                para_format = para.paragraph_format #获取当前段落的格式设置对象，并将其赋值给 para_format
                para_format.alignment = alignment   #设置段落的对齐方式
                para_format.left_indent = Pt(left_indent)   #段落的左缩进和右缩进
                para_format.right_indent = Pt(right_indent)
                para_format.space_before = Pt(space_before) #段落前和段落后的间距
                para_format.space_after = Pt(space_after)
                para_format.line_spacing = Pt(line_spacing) #设置段落的行间距
                para_format.first_line_indent = Pt(first_line_indent)   #设置段落的首行缩进,为啥需要它，呵呵
        except Exception as e:
            raise Exception(f"设置段落格式时出现错误: {e}")

    def paragraph_set(
            self,
            paragraph,
            font_name_ch="方正小标宋简体",
            font_name_west="Times New Roman",
            font_size=22,
            para_alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent=2,
            left_indent=0,
            right_indent=0,
            space_before=0,
            space_after=0
    ):
        """
        设置单个段落的格式，包括字体、对齐方式、缩进、间距等。

        参数:
        paragraph (docx.text.Paragraph): 要设置格式的段落对象
        font_name_ch (str, 可选): 中文字体名称，默认值为"方正小标宋简体"
        font_name_west (str, 可选): 西文字体名称，默认值为"Times New Roman"
        font_size (float, 可选): 字体大小，默认值为22
        para_alignment (WD_ALIGN_PARAGRAPH, 可选): 段落对齐方式，默认值为居中（WD_ALIGN_PARAGRAPH.CENTER）
        first_indent (float, 可选): 首行缩进，单位为字符数（根据字体大小换算为磅），默认值为2
        left_indent (float, 可选): 左缩进，单位为磅，默认值为0
        right_indent (float, 可选): 右缩进，单位为磅，默认值为0
        space_before (float, 可选): 段前间距，单位为磅，默认值为0
        space_after (float, 可选): 段后间距，单位为磅，默认值为0

        返回:
        None

        异常处理:
        - 若在设置段落字体、对齐等格式过程中出现底层库相关问题（如字体不支持等），抛出异常方便排查。
        """
        try:
            paragraph.alignment = para_alignment # 默认设置段落居中
            paragraph_format = paragraph.paragraph_format   #设置行间距
            paragraph.line_space_rule = WD_LINE_SPACING.EXACTLY     #固定值
            paragraph_format.line_spacing = Pt(28)
            paragraph_format.first_line_indent = Pt(font_size * first_indent)
            # 406400代表两字符，先在word上设置好，再用程序反向查找    document.paragraphs[1]. paragraph_format.first_line_indent
            paragraph_format.left_indent = Pt(left_indent)  #设置段落缩进
            paragraph_format.right_indent = Pt(right_indent)
            paragraph_format.space_before = Pt(space_before)#设置段落间距
            paragraph_format.space_after = Pt(space_after)

            for run in paragraph.runs:  # 设置中文字体
                run.font.name = font_name_ch
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_ch)
                run.font.name = font_name_west  # 设置西文字体
                run.font.size = Pt(font_size)   #12:小四，18：小二，22：二号  16：三号
        except Exception as e:
            raise Exception(f"设置段落 {paragraph.text} 格式时出现错误: {e}")

    def paragraphs_set_all(self):
        """
        设置文档所有段落为正文格式（仿宋_GB2312字体，两端对齐等）。

        返回:
        None

        异常处理:
        - 若在遍历段落设置格式过程中出现异常（如内存不足等极端情况或者底层库错误），向上抛出异常。
        """
        try:
            for paragraph in self.document.paragraphs:
                self.paragraph_set(
                    paragraph,
                    font_name_ch="仿宋_GB2312",
                    font_name_west="Times New Roman",
                    font_size=16,
                    para_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_indent=2
                )
        except Exception as e:
            raise Exception(f"设置所有段落格式时出现错误: {e}")

    def paragraphs_set_hides_first(self, lines=[0]):
        """
        设置文档的标题行格式（方正小标宋简体字体，居中对齐等）。

        参数:
        lines (list, 可选): 标题所在的行号列表，默认值为[0]

        返回:
        None

        异常处理:
        - 若在设置标题行格式时出现如索引越界（当传入的行号不合理时）、字体设置失败等问题，抛出异常。
        """
        try:
            for line in lines:
                self.paragraph_set(
                    self.document.paragraphs[line],
                    font_name_ch="方正小标宋简体",
                    font_name_west="Times New Roman",
                    font_size=22,
                    para_alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_indent=0
                )
        except Exception as e:
            raise Exception(f"设置标题行格式时出现错误: {e}")

    def paragraphs_set_date_and_name_under_title(self, lines=[1]):
        """
        设置标题下的日期、姓名等行的格式（楷体_GB2312字体，居中对齐等）。

        参数:
        lines (list, 可选): 对应行号列表，默认值为[1]

        返回:
        None

        异常处理:
        - 类似其他设置行格式的函数，若出现格式设置相关的异常，抛出异常便于排查问题。
        """
        try:
            for line in lines:
                self.paragraph_set(
                    self.document.paragraphs[line],
                    font_name_ch="楷体_GB2312",
                    font_name_west="Times New Roman",
                    font_size=16,
                    para_alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_indent=0
                )
        except Exception as e:
            raise Exception(f"设置标题下相关行格式时出现错误: {e}")

    def paragraphs_set_inscribe(self, lines=[2]):
        """
        设置主送机关、落款单位和时间等行的格式（仿宋_GB2312字体，对应对齐方式等）。

        参数:
        lines (list, 可选): 对应行号列表，默认值为[2],第三行

        返回:
        None

        异常处理:
        - 若在设置这些关键行格式时出现如对齐方式设置失败、字体应用异常等情况，抛出异常。
        """
        try:
            for line in lines:
                self.paragraph_set(
                    self.document.paragraphs[line],
                    font_name_ch="仿宋_GB2312",
                    font_name_west="Times New Roman",
                    font_size=16,
                    para_alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_indent=0
                )

            self.paragraph_set(
                self.document.paragraphs[-2],
                font_name_ch="仿宋_GB2312",
                font_name_west="Times New Roman",
                font_size=16,
                para_alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                first_indent=0
            )

            self.paragraph_set(
                self.document.paragraphs[-1],
                font_name_ch="仿宋_GB2312",
                font_name_west="Times New Roman",
                font_size=16,
                para_alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                first_indent=4
            )
            self.document.paragraphs[-1].paragraph_format.right_indent = Pt(16 * 4)
        except Exception as e:
            raise Exception(f"设置落款相关行格式时出现错误: {e}")

    def paragraphs_set_hides_second(self):
        """
        设置文档二级标题格式（黑体字体，两端对齐等）。

        返回:
        None

        异常处理:
        - 若在设置二级标题格式时出现如格式不匹配、字体应用问题等异常，抛出异常便于排查。
        """
        try:
            for paragraph in self.document.paragraphs:
                string = paragraph.text
                if self.is_title_second(string):
                    self.paragraph_set(
                        paragraph,
                        font_name_ch="黑体",
                        font_name_west="Times New Roman",
                        font_size=16,
                        para_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                    )
        except Exception as e:
            raise Exception(f"设置二级标题格式时出现错误: {e}")

    def paragraphs_set_hides_third(self):
        """
        设置文档三级标题格式（楷体_GB2312与仿宋_GB2312字体结合等）。

        返回:
        None

        异常处理:
        - 若在设置三级标题格式时出现如文本分割、字体设置等环节的异常，抛出异常以便处理。
        """
        try:
            for paragraph in self.document.paragraphs:
                text = paragraph.text
                # 查找第一个句号的位置,但是也存在没有结尾的句号，只有一句话的情况，如何解决。
                index = text.find('。')
                if index!= -1:
                    part1 = text[:index + 1]
                    part2 = text[index + 1:]
                    pattern = r'^（[一二三四五六七八九十]+）(.*?)。'
                    match = re.match(pattern, part1)
                    if match:
                        run1 = paragraph.add_run(part1)
                        run2 = paragraph.add_run(part2)

                        run1.font.name = '楷体_GB2312'
                        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                        run1.font.name = "Times New Roman"
                        run1.font.size = Pt(16)

                        if run2:
                            run2.font.name = '仿宋_GB2312'
                            run2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                            run2.font.name = "Times New Roman"
                            run2.font.size = Pt(16)

                        paragraph.runs[0].clear()
                else:
                    pattern = r'^（[一二三四五六七八九十]+）(.*?)'  #没用到它哈
                    match = re.match(pattern, text)    #没用到它哈
                    if self.is_title_third(text):
                        self.paragraph_set(
                            paragraph,
                            font_name_ch="楷体_GB2312",
                            font_name_west="Times New Roman",
                            font_size=16,
                            para_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                        )
        except Exception as e:
            raise Exception(f"设置三级标题格式时出现错误: {e}")

    def is_title_second(self, string):
        """
        判断字符串是否符合二级标题的格式要求。

        参数:
        string (str): 要判断的字符串

        返回:
        bool: 如果符合二级标题格式要求，返回True，否则返回False
        """
        pattern = r'^[一二三四五六七八九十]+、'
        match = re.match(pattern, string)
        return bool(match)

    def is_title_third(self, string):
        """
        判断字符串是否符合三级标题的格式要求。

        参数:
        string (str): 要判断的字符串

        返回:
        bool: 如果符合三级标题格式要求，返回True，否则返回False
        """
        pattern = r'^（[一二三四五六七八九十]+）'
        match = re.match(pattern, string)
        return bool(match)

    def create_custom_style(self, style_name, font_name_ch, font_name_west, font_size,
                            para_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            left_indent=0, right_indent=0, space_before=0, space_after=0,
                            line_spacing=28, first_line_indent=0):
        """
        创建一个自定义的段落样式。

        参数:
        style_name (str): 样式名称，用于后续应用该样式
        font_name_ch (str): 中文字体名称
        font_name_west (str): 西文字体名称
        font_size (float): 字体大小
        para_alignment (WD_ALIGN_PARAGRAPH, 可选): 段落对齐方式，默认值为两端对齐
        left_indent (float, 可选): 左缩进，单位为磅，默认值为0
        right_indent (float, 可选): 右缩进，单位为磅，默认值为0
        space_before (float, 可选): 段前间距，单位为磅，默认值为0
        space_after (float, 可选): 段后间距，单位为磅，默认值为0
        line_spacing (float, 可选): 行间距，单位为磅，默认值为28
        first_line_indent (float, 可选): 首行缩进，单位为磅

        返回:
        None
        """
        style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = font_name_ch
        style.font.size = Pt(font_size)
        style.font.name = font_name_west
        style.paragraph_format.alignment = para_alignment
        style.paragraph_format.left_indent = Pt(left_indent)
        style.paragraph_format.right_indent = Pt(right_indent)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.line_spacing = Pt(line_spacing)
        style.paragraph_format.first_line_indent = Pt(first_line_indent)

    def apply_custom_style(self, paragraph, style_name):
        """
       应用自定义样式到指定段落

        参数:
        paragraph (docx.text.Paragraph): 要应用样式的段落
        style_name (str): 样式名称

        返回:
        None
        """
        style = self.document.styles[style_name]
        paragraph.style = style

    def set_list_format(self, paragraph, bullet_type='bullet', bullet_char='•', indent=0):
        """
        设置段落列表格式

        参数:
        paragraph (docx.text.Paragraph): 要设置格式的段落
        bullet_type (str, optional): 'bullet' 表示无序列表, 'number' 表示有序列表
        bullet_char (str, optional): 子弹字符
        indent (float, optional): 缩进量

        返回:
        None
        """
        if bullet_type == 'bullet':
            paragraph.style = self.document.styles['List Bullet']
            paragraph.runs[0].text = bullet_char
            paragraph.paragraph_format.left_indent = Pt(indent)
        elif bullet_type == 'number':
            paragraph.style = self.document.styles['List Number']
            paragraph.paragraph_format.left_indent = Pt(indent)

    def set_table_format(self, table, cell_align=WD_ALIGN_PARAGRAPH.CENTER,
                         border_color=RGBColor(0, 0, 0),
                         header_font_size=12, header_font_name='Arial'):
        """
        设置表格格式

        参数:
        table (docx.table.Table): 表格对象
        cell_align (WD_ALIGN_PARAGRAPH, optional): 单元格对齐方式
        border_color (RGBColor, optional): 边框颜色
        header_font_size (int, optional): 表头字体大小
        header_font_name (str, optional): 表头字体

        返回:
        None
        """
        for row in table.rows:
            for cell in row.cells:
                cell.paragraph_format.alignment = cell_align
                if cell.style.name == 'TableHeader':
                    cell.paragraph_format.font.size = Pt(header_font_size)
                    cell.paragraph_format.font.name = header_font
        for cell in table.rows[0].cells:
            cell.style = self.document.styles['TableHeader']
        for row in table.rows:
            for cell in row.cells:
                for border in cell._element.findall('.//w:tcBorders'):
                    border.set(qn('w:color'), border_color.rgb)

    def document_to_offical_format(self, hides_first=[0]):
        # 将源文件转为新的文件，按说也可以不转，直接处置即可

        # 设置页面
        self.pages_set(page_width=21, page_height=29.7, orientation=WD_ORIENT.PORTRAIT)
        # 设置页边距，可以从参数文件中读取，不用修改函数，尝试使用JSON
        self.sections_set(left=2.8, right=2.6, top=3.7, bottom=3.5)
        # 设置段落边距等
        self.para_set_indent()

        # 设置正文，全部先设置为正文格式，再逐个调整各级标题
        self.paragraphs_set_all()
        # 设置各级标题
        self.paragraphs_set_hides_first( lines=hides_first)
        self.paragraphs_set_hides_second()
        self.paragraphs_set_hides_third()
        # 设置主送机关和落款单位、时间
        self.paragraphs_set_inscribe()
        # 添加页码，设置，如果页面超过2页，则添加，另外，如果
        # self.InsertPageNumber()

        # 保存为新的文件
        return self.document
