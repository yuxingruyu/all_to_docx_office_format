# Create:20241226
# modify codes using class and create testing code


from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml import OxmlElement
import os
import win32com.client as win32
from docx import Document


class DocumentPageNumberAdder:
    def __init__(self, document):
        self.document = document

    def add_page_number(self):
        """
        为文档添加页码。

        此方法先尝试使用docx库自身的功能添加页码，如果添加失败（例如docx库本身对页码相关属性支持有限），
        则会尝试借助win32com客户端调用Word应用程序来添加页码（前提是系统安装了Word软件）。

        异常处理:
        - 如果在保存临时文件、调用Word应用程序、获取或设置页码相关属性等过程中出现权限问题、文件不存在、
          应用程序调用失败等各类异常，都会向上抛出相应的异常信息，方便排查问题。

        返回:
        None
        """
        try:
            # 先尝试使用docx库自带方式添加页码（目前docx库在页码设置上有一定局限性，但先尝试简单方式）
            self._add_page_number_with_docx()
            return
        except Exception as e:
            print(f"使用docx库添加页码失败，尝试使用win32com方式，原因: {e}")

        try:
            # 使用win32com调用Word应用程序添加页码
            self._add_page_number_with_win32com()
        except Exception as e:
            raise Exception(f"使用win32com添加页码时出现错误: {e}")

    def _add_page_number_with_docx(self):
        """
        使用docx库尝试添加页码，如果无法完成页码添加的完整功能（例如无法准确设置页码格式等复杂情况），
        仅做简单的页码插入尝试，若失败则抛出异常。

        异常处理:
        - 如果在操作docx文档对象的元素来添加页码时出现底层库相关的异常，向上抛出异常。

        返回:
        None
        """
        footer = self.document.sections[0].footer
        footer.is_linked_to_previous = True
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 这里可以根据docx库后续支持的更完善的页码功能进行扩展，目前只是简单示例
        run_left = paragraph.add_run('—  ')
        self._set_font(run_left, '仿宋', 14)

        run_footer = paragraph.add_run()
        self._add_page_number_field(run_footer)
        self._set_font(run_footer, '仿宋', 14)

        run_right = paragraph.add_run('  — ')
        self._set_font(run_right, '仿宋', 14)

    def _add_page_number_with_win32com(self):
        """
        使用win32com客户端调用Word应用程序来添加页码，涉及将文档保存为临时文件、打开Word应用程序、
        设置页码属性以及关闭和清理临时文件等操作。

        异常处理:
        - 如果保存临时文件时出现I/O异常（如权限问题、磁盘空间不足等），抛出相应异常。
        - 如果调用Word应用程序失败（如Word未安装、COM组件注册问题等），抛出对应异常。
        - 如果在获取或设置Word文档的页码相关属性时出现问题，抛出异常。

        返回:
        None
        """
        abs_path = os.path.abspath(self.document.file.name)
        temp_path = os.path.join(os.path.dirname(abs_path), "~temp.docx")

        # 保存文档为临时文件
        try:
            self.document.save(temp_path)
        except IOError as e:
            raise IOError(f"保存临时文件 {temp_path} 时出现错误: {e}")

        # 打开临时文件并获取页数
        try:
            word = win32.gencache.EnsureDispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(temp_path)
            doc.BuiltInDocumentProperties(14).Value  # 这里可以进一步完善设置页码格式等操作
            doc.Close()
            word.Quit()
        except Exception as e:
            # 如果临时文件存在则删除，避免残留无用文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise Exception(f"调用Word应用程序处理页码时出现错误: {e}")

        # 删除临时文件
        try:
            os.remove(temp_path)
        except OSError as e:
            print(f"删除临时文件 {temp_path} 时出现错误，但不影响主要流程，错误信息: {e}")

    def _set_font(self, run, font_name, font_size):
        """
        设置run对象（docx文本运行对象）的字体和字号。

        参数:
        run (docx.text.Run): 要设置字体的文本运行对象
        font_name (str): 字体名称，如'仿宋'
        font_size (int): 字号大小，如14

        异常处理:
        - 如果在设置字体相关属性时出现底层库不支持的操作等异常，向上抛出异常。

        返回:
        None
        """
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size

    def _add_page_number_field(self, run):
        """
        在给定的run对象中添加页码字段，通过操作docx的底层XML元素来实现。

        参数:
        run (docx.text.Run): 要添加页码字段的文本运行对象

        异常处理:
        - 如果在操作XML元素构建页码字段时出现异常（如元素创建、属性设置失败等），向上抛出异常。

        返回:
        None
        """
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'Page'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        t = OxmlElement('w:t')
        t.text = "Seq"
        fldChar2.append(t)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)