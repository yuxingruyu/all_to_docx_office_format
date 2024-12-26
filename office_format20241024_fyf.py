#!/usr/bin/env python
# coding: utf-8

# In[25]:


# Generating_document_format_yuxingruyu20240730
#生成几个独立的函数
# 1.标准公文，一行标题，主送、落款单位和日期# 
# 2.标准公文，2行标题，主送、落款单位和日期
# 3.普通公文，1行标题， 无主送、落款单位和日期
#4.条例，个别需要修改的地方，如加粗第几章，第几条，居中的地方
#5.缩小版，打印较多的材料用

#Modify:20241217
# 生成可出售版本，带广告
# 免费版本：当前文件夹下，所有docx、txt文件，可以在原基础上生成带主送落款和只有标题与正文的2种，固定格式。
#会员版本，增加定制格式功能，增加子目录修改公文格式功能
#对标网络售卖的其他软件
#Modify:20241024

#使用pycharm,单独使用main.py，生成exe文件

#Modify:20240802
#增加题目下面可加可选的1-2行楷体居中的日期、姓名（领导讲话、发言材料）
#Modify:20240801
#增加题目下面可加可选的1-2行楷体居中的日期、姓名（领导讲话、发言材料）
# 2.标准公文，2行标题，主送、落款单位和日期
# 3.普通公文，1行标题， 无主送、落款单位和日期

#Modify:20240731
#设置为其他程序可调用的
#终于有了自己的成功的库文件
#实现了多类型的参数，如允许参数为docx\txt\str三种格式的一种
#增加二级标题允许以一行的方式存在，注意检查方法和顺序,成功
#完善了部分错误检查，增加容错功能
#Modify:20240730
#把重要的功能、可调用的函数记录下来，时间长了，也能看明白
#把原来特例的函数通用化，加上参数（或关键字参数）
#Modify:20240711
#使用函数，大改

#modify:20230814，设置为一个函数
#增加识别txt文件，并先全部转换为docx文件
#modify:20230706增加生成的文件名上增加日期
#在run中匹配公文二级标题。使用.text
#？？？页面下端会有空白行，原因不详。
#Modify:20230608增加了插入页码和设定A4纸的功能
#迷信于AI和微信搜索，始终没有找到插入页码的方法。通过抖音搜索到了，再查百度，也能找到。
#谁有用，就用谁，不行就换一家。不迷恋，不执著。
#Create:20230607
#插入页码的功能没有实现
#三级标题，再根据实际变化一下。
#标题与第一段落之间要有空行
#已经修正：标题和主送机关不能到最前面顶格。
#使用python-docx实现将word文件调整为公文格式(或指定的格式)
#比较稳妥的方式是新生成一个word,不受原来的干扰，而导入的文件，只读取内容即可，读作文本


# In[ ]:





# In[26]:


# import docx
from docx import Document

from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor

from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement  #这个对不对
from docx.oxml import OxmlElement

from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_ORIENT, WD_SECTION

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE

import os
import re

import shutil

import datetime


from pathlib import Path





# In[ ]:





# In[27]:


#从docx文档中读取全部文本
def get_text(file_name):
    '''从docx文档中读取全部文本'''
    doc = docx.Document(file_name)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


# def docx_old_to_docx_new(docx_path_in):
#     '''
#     #导入需要修改的文件的文本内容     
#     #生成一个新的docx文档，防止被原文档的设置干扰
#     '''
#     doc_in_text = get_text(docx_path_in)

    
#     document = docx.Document()
#     #将导入的文件文本，分成各行，分别生成不同的段落。不分的话，是在同一个段落中。
    
#     list_text = doc_in_text.split('\n')#不小心实现了删除全部空行
#     for text in list_text:
#         document.add_paragraph(text)
    
#     return document

# import docx
# import os

def docx_old_to_docx_new(input_data):
    '''
    # 导入需要修改的文件的文本内容     
    # 生成一个新的 docx 文档，防止被原文档的设置干扰
    '''
    if isinstance(input_data, str):
        if os.path.isfile(input_data) and input_data.endswith('.docx'):
            doc_in_text = get_text(input_data)
        else:
            doc_in_text = input_data
    elif isinstance(input_data, docx.document.Document):
        doc_in_text = '\n'.join([p.text for p in input_data.paragraphs])
    else:
        raise ValueError("输入参数类型不正确，应为 str（docx 文件路径）或 docx.document.Document 类型")

    document = docx.Document()
    # 将导入的文件文本，分成各行，分别生成不同的段落。不分的话，是在同一个段落中。
    list_text = doc_in_text.split('\n')  # 不小心实现了删除全部空行
    for text in list_text:
        if text.strip():  # 增加判断，去除空行
            document.add_paragraph(text)

    return document

def pages_set(document,page_width = 21 ,page_height =29.7 ,orientation = WD_ORIENT.PORTRAIT):
    '''
    #设置纸张大小,默认为A4，，使用单位为厘米，纵向'''
    # 创建一个新的section对象
    section = document.sections[-1]  #为什么只设置最后一个节的就可以

    # 设置纸张大小为A4
    section.page_width = Cm(page_width)
    section.page_height = Cm(page_height)

    # 设置页面方向为纵向
    section.orientation = orientation

    
#设置页边距
def sections_set(document, left = 2.8, right = 2.6, top = 3.7, bottom = 3.5):
    sections = document.sections #为什么设置所有的节的
    for section in sections:
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)

        
        
        
#匹配二级标题（公文的一级标题）
def is_title_second(string):
    pattern = r'^[一二三四五六七八九十]+、'
    match = re.match(pattern, string)
    if match:
        return True
    else:
        return False

    
#匹配三级标题（公文的二级标题，与实际还不同，要修改）
def is_title_third(string):
    pattern = r'^（[一二三四五六七八九十]+）'
    match = re.match(pattern, string)
    if match:
        return True
    else:
        return False
    
#通用匹配，尽量函数化，减少独立的类似的函数
def match_str(string, pattern):
#     pattern = r'^（[一二三四五六七八九十]+）'
    '''匹配文字内容是否符合pattern要求，如各级标题'''
    match = re.match(pattern, string)
    if match:
        return True
    else:
        return False
    
#设置段落格式，左右上下缩进为0，行距28磅

def para_set_indent(
    document,
    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY,
    left_indent = 0,
    right_indent = 0,
    space_before = 0,
    space_after = 0,
    line_spacing = 28,
    first_line_indent = 28
):
    '''设置段落间距，默认使用公文的'''
    for para in document.paragraphs: #遍历 document 中的所有段落
        para_format = para.paragraph_format #获取当前段落的格式设置对象，并将其赋值给 para_format
        para_format.alignment = alignment #设置段落的对齐方式
        para_format.left_indent = Pt(left_indent) #段落的左缩进和右缩进
        para_format.right_indent = Pt(right_indent)
        para_format.space_before = Pt(space_before)#段落前和段落后的间距
        para_format.space_after = Pt(space_after)
        para_format.line_spacing = Pt(line_spacing)#设置段落的行间距
        para_format.first_line_indent = Pt(first_line_indent) #设置段落的首行缩进,为啥需要它，呵呵
        
        
#设置段落，行间距28，首行缩进2字符，段前后为0，两端对齐,中文字体，西文字体
#"方正小标宋简体""黑体""楷体_GB2312""仿宋_GB2312"

def paragraph_set(
    paragraph, 
    font_name_ch = "方正小标宋简体", 
    font_name_west = "Times New Roman", 
    font_size = 22 ,
    para_alignment = WD_ALIGN_PARAGRAPH.CENTER, 
    first_indent = 2,
    left_indent = 0,
    right_indent = 0,
    space_before = 0,
    space_after = 0
):
    # 设置段落居中
    paragraph.alignment = para_alignment
    
    #设置行间距
    paragraph_format = paragraph.paragraph_format
    paragraph.line_space_rule = WD_LINE_SPACING.EXACTLY #固定值
    paragraph_format.line_spacing = Pt(28)
    #paragraph_format.first_line_indent = Pt(28)
    #paragraph_format.first_line_indent = 406400   #406400代表两字符，先在word上设置好，再用程序反向查找    document.paragraphs[1]. paragraph_format.first_line_indent
    paragraph_format.first_line_indent = Pt(font_size * first_indent)  #？？？感觉是字体大小的两倍，未验证。
    
    #设置段落缩进
    paragraph_format.left_indent = Pt(left_indent)
    paragraph_format.right_indent = Pt(right_indent)
    #设置段落间距
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    
    
    # 设置中文字体
    for run in paragraph.runs:
        run.font.name = font_name_ch
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_ch)

        # 设置西文字体
        run.font.name = font_name_west
        run.font.size = Pt(font_size)
        #12:小四，18：小二，22：二号  16：三号    

        
# 设置正文
def paragraphs_set_all(document):
    #先设置正文格式
    for paragraph in document.paragraphs:
        paragraph_set(paragraph,font_name_ch = "仿宋_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent = 2)


#设置各级标题
def paragraphs_set_hides_first(document,lines = [0]):
#     lines = [0],取第0行
    #设置标题,有时侯需要设置2行
    for line in lines:
        paragraph_set(document.paragraphs[line],font_name_ch = "方正小标宋简体", font_name_west = "Times New Roman", font_size = 22, para_alignment = WD_ALIGN_PARAGRAPH.CENTER, first_indent = 0 )
    #document.paragraphs[0].paragraph_format.left_indent = Pt(0)
    #document.paragraphs[0].paragraph_format.first_line_indent = Pt(0)

#设置各级标题
def paragraphs_set_date_and_name_under_title(document,lines = [1]):
    #设置标题,有时侯需要设置2行
    for line in lines:
        paragraph_set(document.paragraphs[line],font_name_ch = "楷体_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.CENTER, first_indent = 0 )
    #document.paragraphs[0].paragraph_format.left_indent = Pt(0)
    #document.paragraphs[0].paragraph_format.first_line_indent = Pt(0)

    #设置主送机关、落款单位和时间
def paragraphs_set_inscribe(document,lines = [2]):
    #设置主送机关,第三行
    for line in lines:
        paragraph_set(document.paragraphs[line],font_name_ch = "仿宋_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.LEFT, first_indent = 0 )
    #document.paragraphs[2].paragraph_format.left_indent = Pt(0)    
    #document.paragraphs[2].paragraph_format.first_line_indent = Pt(0)
        #设置落款单位
    paragraph_set(document.paragraphs[-2],font_name_ch = "仿宋_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.RIGHT, first_indent = 0 )
        #设置落款时间，最后面空4格

    paragraph_set(document.paragraphs[-1],font_name_ch = "仿宋_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.RIGHT, first_indent = 4 )
    document.paragraphs[-1].paragraph_format.right_indent = Pt( 16 * 4)

    
    
#设置二级标题
def paragraphs_set_hides_second(document):
    for paragraph in document.paragraphs:
        string = paragraph.text
        if is_title_second(string):
            paragraph_set(paragraph,font_name_ch = "黑体", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY)

            
#pattern = r'^（[一二三四五六七八九十]+）(.*)。'


#设置三级标题，只第一句。
def paragraphs_set_hides_third(document):
    for paragraph in document.paragraphs:
        text = paragraph.text
        # 查找第一个句号的位置,但是也存在没有结尾的句号，只有一句话的情况，如何解决。
        index = text.find('。')

        if index != -1:
            # 将文本分成两个部分
            part1 = text[:index+1]
            part2 = text[index+1:]

            pattern = r'^（[一二三四五六七八九十]+）(.*?)。'
            '''在这个模式中：
            ^ 依旧表示匹配字符串的开头。
            （[一二三四五六七八九十]+） ：
            （ 和 ） 是汉字的左括号和右括号。
            [一二三四五六七八九十]+ 表示匹配由“一”到“十”这些汉字组成的一个或多个连续的字符序列。
            (.*?) ：
            .* 表示匹配任意字符（除了换行符）零次或多次。
            ? 表示非贪婪模式，即尽可能少地匹配字符，以确保不会过度匹配后面的内容。
            。 匹配汉字的句号“。”
            这个正则表达式整体用于匹配以汉字左括号开头，接着是由“一”到“十”这些汉字组成的一个或多个连续的字符序列，再跟着汉字右括号，然后是任意内容，最后以汉字句号结束的字符串。
            例如，它可以匹配 “（一） 这是一些内容。” 、“（八） 更多的内容。” 等字符串。'''
            
            match = re.match(pattern, part1)



            if match:
                # 创建两个run
                run1 = paragraph.add_run(part1)
                run2 = paragraph.add_run(part2)

                # 设置run的格式和样式
                run1.font.name = '楷体_GB2312'
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')

                # 设置西文字体
                run1.font.name = "Times New Roman"
                run1.font.size = Pt(16)

                if run2:
                    # 设置run的格式和样式
                    run2.font.name = '仿宋_GB2312'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

                    # 设置西文字体
                    run2.font.name = "Times New Roman"
                    run2.font.size = Pt(16)

                paragraph.runs[0].clear()
                
        #如果没有句号，再判断是否只有一句话
        else:
            pattern = r'^（[一二三四五六七八九十]+）(.*?)'
            match = re.match(pattern, text)
            if is_title_third(text):
                paragraph_set(paragraph,font_name_ch = "楷体_GB2312", font_name_west = "Times New Roman", font_size = 16, para_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY)



            


# In[28]:


import docx

def is_single_sentence(paragraph_text):
    """
    此函数用于判断一个段落是否只有一句话

    参数：
    paragraph_text (str): 要判断的段落文本

    返回：
    bool: 如果段落符合只有一句话的条件，返回 True，否则返回 False
    """
    # 去除段落两端的空白
    paragraph_text = paragraph_text.strip()
    if not paragraph_text:  # 如果段落为空，返回 False
        return False

    # 按常见的标点符号分割段落文本
    sentences = re.split(r'[。!?]', paragraph_text)
    # 如果分割后只有一个部分，并且这个部分的结尾不是逗号或者没有标点且不是因为换行导致的分割，那么认为是一句话
    if len(sentences) == 1 and not (paragraph_text.endswith(',') or (len(paragraph_text) > 0 and paragraph_text[-1] not in ['.', '!', '?', ','] and not paragraph_text.endswith('\n'))):
        return True
    else:
        return False


# In[29]:


#在python docx中实现word插入页码

#主要使用的是页脚功能。

#另外，用到了docx的oxml命令。
#https://baijiahao.baidu.com/s?id=1665454009794833226&wfr=spider&for=pc
#https://stackoverflow.com/questions/50776715/setting-pgnumtype-property-in-python-docx-is-without-effect

def AddFooterNumber(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
def InsertPageNumber(Doc):
    footer = Doc.sections[0].footer # 获取第一个节的页脚
    footer.is_linked_to_previous = True  #编号续前一节
    paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页脚居中对齐
    
    run_left = paragraph.add_run()
    run_left.text = '—  '
    font = run_left.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗
    
    run_footer=paragraph.add_run() # 添加页脚内容
    AddFooterNumber(run_footer)
    font = run_footer.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_footer._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗
    
    run_right = paragraph.add_run()
    run_right.text = '  — '
    font = run_right.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_right._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗

#EndDef


# In[30]:


def delete_paragraphs_null_lines(document):    
    '''删除docx文件的空行'''
    # 打开docx文件
    doc = document

    # 创建新文档
    new_doc = Document()

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        # 如果段落不是空行，则复制到新文档中
        if paragraph.text.strip():
            new_doc.add_paragraph(paragraph.text)

    # 保存修改后的文件
    #new_doc.save('modified_file.docx')
    return new_doc


# In[31]:


#将txt文件转换为docx文件
def txt_to_docx(txt_file, docx_file):
    # 打开txt文件
    with open(txt_file, 'r', encoding = 'utf-8') as f:
        content = f.read()

    # 将txt内容按照换行符进行分割,确保换行符不变。
    lines = content.split('\n')

    # 创建一个新的Word文档
    doc = Document()

    # 将每一行文本添加到docx文档中
    for line in lines:
        doc.add_paragraph(line)


    # 保存为docx文件
    doc.save(docx_file)
    return doc
    
#将指定文件夹中的txt文件转换为docx文件
def txt_to_docx_all(folder_path_in):
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.txt'):
            txt_path_in = os.path.join(folder_path_in, file_name_in)
            #docxs.append(docx_path_in)
            docx_path_out = txt_path_in[:-3] + 'docx'
            
            txt_to_docx(txt_path_in, docx_path_out)


# In[32]:


# 大标题占用一行或2行
def docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1]):
    #将源文件转为新的文件，按说也可以不转，直接处置即可
    document = docx_old_to_docx_new(docx_path_in)
    #设置页面
    pages_set(document)
    #设置页边距，可以从参数文件中读取，不用修改函数，尝试使用JSON
    sections_set(document, left = 2.8, right = 2.6, top = 3.7, bottom = 3.5)
    #设置段落边距等
    para_set_indent(document)
    
    #设置正文，全部先设置为正文格式，再逐个调整各级标题
    paragraphs_set_all(document)
    #设置各级标题
    paragraphs_set_hides_first(document,lines = hides_first)
    paragraphs_set_hides_second(document)
    paragraphs_set_hides_third(document)
    #设置主送机关和落款单位、时间
    paragraphs_set_inscribe(document)
    #添加页码，设置，如果页面超过2页，则添加，另外，如果
    InsertPageNumber(document)
    
    #保存为新的文件
    document.save(docx_path_out)
    


# In[33]:


import docx
import json

def file_to_docx(input_file, docx_path_out, hides_first=[0, 1],have_inscribe = True,date_and_name_under_title_flag = False,date_and_name_under_title_numbers = [1]):
    """
    根据输入文件的类型进行处理，如果是docx文件则直接处理，
    如果是txt文件则转换为docx后处理，如果是字符串则直接处理

    参数：
    input_file (str): 输入的文件路径或字符串
    docx_path_out (str): 输出的docx文件路径
    hides_first (list, 可选): 一些控制参数，默认为 [0, 1]
    默认有主送、落款
    返回：
    None
    """
    if input_file.endswith('.txt'):
        # 将txt文件转换为docx格式并处理
        doc = txt_to_docx_converted(input_file)
    elif input_file.endswith('.docx'):
        # 处理输入为docx文件的情况
        doc = docx.Document(input_file)
    elif isinstance(input_file, str):
        # 处理输入为字符串的情况
        doc = str_to_docx(input_file)
    else:
        # 处理输入为docx文件的情况
        doc = docx.Document(input_file)

    doc = docx_old_to_docx_new(doc)
    # 设置页面
    pages_set(doc)
    # 设置页边距，可以从参数文件中读取，不用修改函数，尝试使用 JSON
    sections_set(doc, left=2.8, right=2.6, top=3.7, bottom=3.5)
    # 设置段落边距等
    para_set_indent(doc)

    # 设置正文，全部先设置为正文格式，再逐个调整各级标题
    paragraphs_set_all(doc)
    # 设置各级标题
    paragraphs_set_hides_first(doc, lines=hides_first)
    paragraphs_set_hides_second(doc)
    paragraphs_set_hides_third(doc)

    #如果有标题下面的日期和姓名，则改之
    if date_and_name_under_title_flag:
        paragraphs_set_date_and_name_under_title(doc,lines = date_and_name_under_title_numbers)

    # 设置主送机关和落款单位、时间,默认为真，执行之
    if have_inscribe: # = True
        paragraphs_set_inscribe(doc)

    # 添加页码，设置，如果页面超过 2 页，则添加，另外，如果
    InsertPageNumber(doc)

    # 保存为新的文件
    doc.save(docx_path_out)
    return doc


# 将txt文件转换为docx文件
def txt_to_docx_converted(txt_file):
    """
    打开txt文件，将内容转换为docx格式

    参数：
    txt_file (str): txt 文件路径

    返回：
    docx.Document: 转换后的 Word 文档对象
    """
    with open(txt_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 将txt内容按照换行符进行分割, 确保换行符不变。
    lines = content.split('\n')

    # 创建一个新的 Word 文档
    doc = docx.Document()

    # 将每一行文本添加到 docx 文档中
    for line in lines:
        doc.add_paragraph(line)

    return doc

import docx

def str_to_docx(input_str):
    """
    将输入字符串转换为docx文档
    有多行的，要分行添加到docx文档的段落中

    参数：
    input_str (str): 输入字符串

    返回：
    docx.Document: 转换后的 Word 文档对象
    """
    doc = docx.Document()
    paragraphs = input_str.split('\n')
    '''当使用 input_str.split('\n') 来分割字符串时，
    如果存在连续的两个 '\n\n' ，
    那么会将这两个连续的换行符视为一个分隔符，
    分割出来的段落列表中，
    在对应位置会是一个空字符串元素。
    '''
    for paragraph in paragraphs:
        if paragraph:
            doc.add_paragraph(paragraph)
    return doc


# In[34]:


# import os
# import shutil

def move_file_to_new_directory(absolute_file_path, new_directory_name):
    '''python,编写一个函数或类，
    参数为一个文件的绝对目录和新目录的名字（当前目录下的新目录），
    查看有无新目录，没有则新建，有则忽略。
    然后将这个文件移动到这个新目录中。
    '''
    current_directory = os.path.dirname(absolute_file_path)
    new_directory_path = os.path.join(current_directory, new_directory_name)
    if not os.path.exists(new_directory_path):
        os.mkdir(new_directory_path)
    shutil.move(absolute_file_path, new_directory_path)
    
    


def copy_file_to_new_directory(absolute_file_path, new_directory_name):
    '''
    python,编写一个函数或类，
    参数为一个文件的绝对目录和新目录的名字（当前目录下的新目录），
    查看有无新目录，没有则新建，有则忽略。然后将这个文件复制到这个新目录中。
    '''
    current_directory = os.path.dirname(absolute_file_path)
    new_directory_path = os.path.join(current_directory, new_directory_name)
    if not os.path.exists(new_directory_path):
        os.mkdir(new_directory_path)
    shutil.copy(absolute_file_path, new_directory_path)


# In[35]:


# import os

def create_directory_if_not_exists(directory_path, folder_name):
    """
    此函数用于检查指定目录中是否存在给定的文件夹，如果不存在则创建

    参数：
    directory_path (str): 要检查的目录路径
    folder_name (str): 要检查和创建的文件夹名称

    返回：
    None
    """
    target_directory = os.path.join(directory_path, folder_name)
    if not os.path.exists(target_directory):
        os.makedirs(target_directory)
    return target_directory


# In[36]:


#docx_to_offical_format(docx_path_in, docx_path_out)


# In[37]:


#批量修改为公文格式，并保存在独立的文件夹中，但复用不方便
def modify_offical_docxs(folder_path_in):
    '''
    将指定文件夹中的txt和docx文件，
    转为标准公文格式（带主送和落款，标题占用2行）
    新文件保存在folder_path_in的兄弟目录“新生成的标准公文格式文件”中
    '''
    date = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    txt_to_docx_all(folder_path_in)
    #docxs = []
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.docx'):
            docx_path_in = os.path.join(folder_path_in, file_name_in)
            #docxs.append(docx_path_in)
#             folder_path = os.path.dirname(folder_path_in) #将folder_path变为局部变量，防止出错,只生成了当前目录，不明所以
            folder_path = str(Path(folder_path_in).parent).replace('\\','/')
            folder_path_out = create_directory_if_not_exists(folder_path,   '生成的标准公文格式文件/')
            docx_path_out = os.path.join(folder_path_out , (date + file_name_in))  #此处有问题，使用了全局变量folder_path，已修改
        
            docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])
            os.startfile(docx_path_out)
            
            
#批量修改为公文格式，直接在原目录生成新的同名文件，增加标准公文格式字样，或其他字样
def modify_offical_docxs_in_current_directory(folder_path_in, folder_relative_path_out = '生成的文件',new_directory_name = '旧',open_new = True):
    # folder_relative_path_out = None 是否可以？
    '''
    将指定文件夹中的txt和docx文件，
    转为标准公文格式（带主送和落款，标题占用2行）
    新文件保存在folder_path_in的子目录folder_relative_path_out中，没有则生成一个
    将原文件保存在子目录new_directory_name中
    '''
    date = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    folder_path_out = create_directory_if_not_exists(folder_path_in,  folder_relative_path_out)

    #不要生成中间文件夹了吧？
    '''
    读取folder_path_in中的txt和docx文件
    将txt转为docx
    将docx们生成带主送和落款的标准公文格式docx
    '''
    # create_directory_if_not_exists(folder_path_in,  new_directory_name)

    txt_to_docx_all(folder_path_in)  #这个要看看保存到了哪里，是同目录下吧？
    #docxs = []
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.docx'):
            docx_path_in = os.path.join(folder_path_in, file_name_in)
            #docxs.append(docx_path_in)
            docx_path_out = os.path.join(folder_path_out , (date + file_name_in))
            # copy_file_to_new_directory(docx_path_in, new_directory_name)
        
            docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])
#             docx_to_offical_format(docx_path_in, docx_path_in,hides_first =  [0,1])  #直接覆盖原文件，可否？
#             os.startfile(docx_path_in) 不需要打开此文件

            print(f'Format {docx_path_in} to {docx_path_out}\n')
            if open_new:
                os.startfile(docx_path_out)


# In[38]:


# import docx

import re

def is_single_sentence(paragraph_text):
    """
    此函数用于判断一个段落是否只有一句话

    参数：
    paragraph_text (str): 要判断的段落文本

    返回：
    bool: 如果段落符合只有一句话的条件，返回 True，否则返回 False
    """
    paragraph_text = paragraph_text.strip()
    if not paragraph_text:
        return False

    # 去除换行符
    paragraph_text = paragraph_text.replace('\n', '')

    # 按句号、感叹号、问号分割
    sentences = re.split(r'[。!?]', paragraph_text)
    sentences = [item for item in sentences if item!= '']  #要删除列表中为空''的部分
    print(sentences)
    # 检查分割后的部分数量以及最后一个部分是否只是标点后的剩余部分（如：“这是一句话。” 分割后为 ["这是一句话", ""] ）
    if len(sentences) == 1 and not sentences[0] == '':
        return True
    else:
        return False


# In[ ]:





# In[39]:


# import os

def find_docx_files(directory):
    '''
    python编写一个函数或类，参数为一个目录，
    查找一个目录中所有以.docx为后缀的文件，
    将全部文件的绝对路径保存到一个列表中，
    并返回。'''
    docx_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                docx_files.append(file_path)
    return docx_files

import os  # 导入 os 模块，用于文件和目录操作

def find_files_by_suffix(directory, suffix, include_subdirectories=False):  # 定义函数 find_files，接受目录、后缀和是否包含子目录的参数
    """
    此函数用于根据给定的目录和后缀（或后缀列表）查找文件，并根据是否包含子目录的参数决定是否遍历子目录，将符合条件的文件的绝对路径保存到列表中返回

    参数：
    directory (str): 要搜索的目录路径
    suffix (str 或 list): 文件的后缀或后缀列表,或后缀元组
    include_subdirectories (bool, 可选): 是否包含子目录，默认为 False

    返回：
    list: 包含符合条件文件的绝对路径的列表
    
    
    """
    file_paths = []  # 创建一个空列表用于存储文件路径
    for root, dirs, files in os.walk(directory):  # 使用 os.walk 遍历目录及其子目录
        """
        root: 当前遍历的目录路径
        dirs: 当前目录下的子目录列表
        files: 当前目录下的文件列表
        """
        if not include_subdirectories and root!= directory:  # 如果不包含子目录且当前目录不是指定目录，则跳过
            continue
        for file in files:  # 遍历当前目录下的文件
            if file.endswith(tuple(suffix)) or (isinstance(suffix, list) and any(file.endswith(suffix_item) for suffix_item in suffix)):  # 判断文件后缀是否符合条件
                """
                当 suffix 是列表时，将其转换为元组传递给 endswith 方法
                
                这里进行后缀的匹配判断。
                `file.endswith(suffix)` 检查文件的后缀是否与给定的单个后缀 `suffix` 相匹配。
                如果 `suffix` 是一个列表，`isinstance(suffix, list) and any(file.endswith(suffix_item) for suffix_item in suffix)` 会检查文件的后缀是否与列表中的任何一个后缀相匹配。
                `any()` 函数用于判断是否存在至少一个满足条件的情况。
                """
                file_paths.append(os.path.abspath(os.path.join(root, file)))  # 将符合条件的文件的绝对路径添加到列表中
                """
                如果文件的后缀符合条件，使用 `os.path.abspath(os.path.join(root, file))` 计算文件的绝对路径，并将其添加到 `file_paths` 列表中。
                `os.path.join(root, file)` 用于组合目录路径 `root` 和文件名 `file` ，形成完整的文件路径。
                `os.path.abspath()` 用于获取该路径的绝对路径。
                """
    return file_paths  # 返回包含文件绝对路径的列表


# In[40]:


import docx
import os
import win32com.client as win32
#可以获取页面，但是依赖于安装了word,慎用
def get_word_page_count(document_path):
    document = docx.Document(document_path)
    abs_path = os.path.abspath(document_path)

    # 保存文档为临时文件
    temp_path = os.path.join(os.path.dirname(abs_path), "~temp.docx")
    document.save(temp_path)

    # 打开临时文件并获取页数
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(temp_path)
    pages = doc.BuiltInDocumentProperties(14).Value
    doc.Close()
    word.Quit()

    # 删除临时文件
    os.remove(temp_path)

    return pages

# document_path = docx_path_out_txt
# page_count = get_word_page_count(document_path)
# print("Word文档的页数：", page_count)
# 没想到获取个word页数这么简单的需求，实现起来还挺难。

# 1.根据网上说法：word文档中不保存总页数，只是在渲染时才根据文字和表格等去摆放到一页上，摆放不下时才展现下一页。所以从docx中取不到总页数

# 2.因此，仅用docx库不行，还需要用win32com。而win32com实际是：pip install pypiwin32   并且一定记得要重启python（重启仍不行可以试试https://blog.csdn.net/weixin_43149311/article/details/120806116  ，我重启IDLE就行了）

# 要获取Word文档的页数，我们可以使用以下步骤：

# 首先，将Word文档保存为临时文件。
# 然后，使用win32com.client模块打开该临时文件。
# 接下来，获取打开的Word应用程序的ActiveDocument。
# 最后，通过BuiltInDocumentProperties属性获取文档的页数。

# 补充：今天执行上述代码出错： 'win32com.gen_py....' has no attribute 'CLSIDToClassMap'

#  解决办法：删除目录<code>C:\Users\你的用户名\AppData\Local\Temp\gen_py\3.X中的缓存文件夹00020905-0000-0000-C000-000000000046x0x8x5即可，重新执行上述代码便不再报错。

 

# 参考：https://blog.51cto.com/u_16175516/6850938

# https://blog.csdn.net/u011840075/article/details/124696371

# https://blog.csdn.net/weixin_43149311/article/details/120806116

# https://blog.51cto.com/mouday/5018378

# https://blog.csdn.net/ericatardicaca/article/details/90721909


# In[ ]:





# In[ ]:





# In[41]:


#获取当前时间
date = datetime.datetime.now().strftime("%Y%m%d")
#当前主文件夹
# folder_path = 'D:/工作自动化/调整公文格式带主送落款'#最后面还有个/，别遗漏
# folder_path = 'D:/工作自动化/生成公文格式20240730'#最后面还有个/，别遗漏

folder_path = Path.cwd()
folder_path = os.path.normpath(os.getcwd()).replace('\\', '/') + '/'
#也可以获取当前文件夹

folder_path_in = folder_path + '/未调整的源文件/'
# /folder_path_in = folder_path / Path( '未调整的源文件')
#文件夹中只保存1个docx文件
# docxs = []
# for file_name_in in os.listdir(folder_path_in):
#     if file_name_in.endswith('.docx'):
#         docx_path_in = os.path.join(folder_path_in, file_name_in)
#         docxs.append(docx_path_in)
#         docx_path_out = os.path.join(folder_path , ( '标准公文格式文件/'+ date + file_name_in))

# print(docx_path_out)


# In[42]:


# 测试用，删除即可
docx_path_out_str = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111str生成的.docx'
docx_path_out_txt = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111txt生成的.docx'
docx_path_out_txt222 = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\222txt生成的.docx'
docx_path_out_docx = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111doc生成的.docx'
# doc = str_to_docx(input_file)
txt_path_in_tmp = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111.txt'
txt_path_in_tmp222 = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\222.txt'
docx_path_in_tmp = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111.docx'
# file_to_docx(input_file, docx_path_out_str, hides_first=[0, 1])
# file_to_docx(docx_path_in_tmp, docx_path_out_docx, hides_first=[0, 1])
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0, 1])
# file_to_docx(input_file, docx_path_out, hides_first=[0, 1],have_inscribe = True)
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0],have_inscribe = False)


# In[43]:


# import pprint
# pprint.pprint(find_files_by_suffix(folder_path_in, 'docx', include_subdirectories=False))


# In[44]:


#将一个docx\txt\srt作为输入，然后生成公文格式文件，单独调用这个最方便
# file_to_docx(path_in, docx_path_out, hides_first=[0, 1])


# In[45]:


#    将指定文件夹中的txt和docx文件，
#     转为标准公文格式（带主送和落款，标题占用2行）
#     新文件保存在folder_path_in的兄弟目录“新生成的标准公文格式文件”中

# modify_offical_docxs(folder_path_in)


# path_tmp = 'D:/code_jupyter/文本转公文格式全自动20241024/kwkw哈哈'
# modify_offical_docxs(path_tmp)

# In[46]:


#os.startfile(docx_path_out)


# In[47]:


# modify_offical_docxs_in_current_directory(folder_path_in, folder_relative_path_out = '生成的文件',new_directory_name = '旧')


# In[48]:


# modify_offical_docxs(folder_path_in)


# In[49]:


#把一个文件转换，调用这个的会比较多，但是不能直接把字符串转换，这样多少会降低效率
# docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])


# In[51]:


#可调用的函数为：
# 1.标准公文，一行标题，主送、落款单位和日期# 
#生成标准公文格式的，带主送和落款，可选首行的行数（1行或2行），后面的可能有bug,主送那里
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0, 1])  

# 2.标准公文，2行标题，主送、落款单位和日期
# 3.普通公文，1行标题， 无主送、落款单位和日期
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0],have_inscribe = False)
#4.条例，个别需要修改的地方，如加粗第几章，第几条，居中的地方
#5.缩小版，打印较多的材料用
#6.领导发言（第二行居中楷体的那种,也可能有日期，即标题下有2行居中楷体）

# file_to_docx(txt_path_in_tmp222, docx_path_out_txt222, hides_first=[0],have_inscribe = False,date_and_name_under_title_flag = True,date_and_name_under_title_numbers = [1])


# In[ ]:


# file_to_docx(
#     txt_path_in_tmp222, 
#     docx_path_out_txt222,
#     hides_first=[0],
#     have_inscribe = False,
#     date_and_name_under_title_flag = True,
#     date_and_name_under_title_numbers = [1]
# )









