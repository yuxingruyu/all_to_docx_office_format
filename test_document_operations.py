# Create:20241226
#modify codes using class and create testing code
from docx import Document, document
from document_converter import DocumentConverter
import os



from document_format_setter import DocumentFormatSetter
from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.oxml.ns import qn

from document_page_number_adder import DocumentPageNumberAdder
from docx import Document

# 测试docx_old_to_docx_new 方法，传入docx文件路径
def test_docx_old_to_docx_new_with_file_path():
    converter = DocumentConverter()
    input_docx_path = "test_input.docx"
    if os.path.exists(input_docx_path):
        new_doc = converter.docx_old_to_docx_new(input_docx_path)
        new_doc.save("test_output_docx_new.docx")
        print("docx_old_to_docx_new with file path test passed!")
    else:
        print(f"文件 {input_docx_path} 不存在，无法进行测试")

# 测试docx_old_to_docx_new 方法，传入已经打开的Document对象
def test_docx_old_to_docx_new_with_document_object():
    # from docx import Document
    converter = DocumentConverter()
    doc = Document()
    doc.add_paragraph("这是一段测试文本")
    new_doc = converter.docx_old_to_docx_new(doc)
    new_doc.save("test_output_docx_new_from_object.docx")
    print("docx_old_to_docx_new with document object test passed!")

# 测试txt_to_docx 方法
def test_txt_to_docx():
    converter = DocumentConverter()
    input_txt_path = "test_input.txt"
    output_docx_path = "test_output.txt_to_docx.docx"
    if os.path.exists(input_txt_path):
        result_doc = converter.txt_to_docx(input_txt_path, output_docx_path)
        print("txt_to_docx test passed!")
    else:
        print(f"文件 {input_txt_path} 不存在，无法进行测试")

# 测试txt_to_docx_all 方法
def test_txt_to_docx_all():
    converter = DocumentConverter()
    input_folder_path = "test_folder"
    if os.path.exists(input_folder_path):
        converter.txt_to_docx_all(input_folder_path)
        print("txt_to_docx_all test passed!")
    else:
        print(f"文件夹 {input_folder_path} 不存在，无法进行测试")

# 测试get_text 方法
def test_get_text():
    converter = DocumentConverter()
    input_docx_path = "test_input.docx"
    if os.path.exists(input_docx_path):
        text_content = converter.get_text(input_docx_path)
        print(f"获取到的文档文本内容（部分展示）: {text_content[:50]}...")
        print("get_text test passed!")
    else:
        print(f"文件 {input_docx_path} 不存在，无法进行测试")

test_docx_old_to_docx_new_with_file_path()
test_docx_old_to_docx_new_with_document_object()
test_txt_to_docx()
test_txt_to_docx_all()
test_txt_to_docx_all()
test_get_text()




# from document_format_setter import DocumentFormatSetter
# from docx import Document
#
# from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
# from docx.enum.section import WD_ORIENTATION, WD_ORIENT, WD_SECTION
# from docx.enum.style import WD_STYLE_TYPE, WD_STYLE
# from docx.shared import Cm, Pt, Inches, RGBColor
# from docx.oxml.ns import qn
# 测试pages_set 方法
def test_pages_set():
    document = Document()
    setter = DocumentFormatSetter(document)
    setter.pages_set(page_width=20, page_height=25, orientation=WD_ORIENTATION.LANDSCAPE)
    document.save("test_pages_set.docx")
    print("pages_set test passed!")

# 测试sections_set 方法
def test_sections_set():
    document = Document()
    setter = DocumentFormatSetter(document)
    setter.sections_set(left=2, right=2, top=3, bottom=3)
    document.save("test_sections_set.docx")
    print("sections_set test passed!")

# 测试para_set_indent 方法
def test_para_set_indent():
    document = Document()
    document.add_paragraph("这是一段测试段落")
    setter = DocumentFormatSetter(document)
    setter.para_set_indent(alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, left_indent=10, right_indent=10,
                           space_before=10, space_after=10, line_spacing=18,
                           first_line_indent=18)
    document.save("test_para_set_indent.docx")
    print("para_set_indent test passed!")

# 测试paragraph_set 方法
def test_paragraph_set():
    document = Document()
    document.add_paragraph("这是一段测试段落")
    setter = DocumentFormatSetter(document)
    setter.paragraph_set(document.paragraphs[0], font_name_ch="宋体", font_name_west="Arial",
                         font_size=14, para_alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, first_indent=1,
                         left_indent=5, right_indent=5, space_before=5, space_after=5)
    document.save("test_paragraph_set.docx")
    print("paragraph_set test passed!")

# 测试paragraphs_set_all 方法
def test_paragraphs_set_all():
    document = Document()
    document.add_paragraph("段落1")
    document.add_paragraph("段落2")
    setter = DocumentFormatSetter(document)
    setter.paragraphs_set_all()
    document.save("test_paragraphs_set_all.docx")
    print("paragraphs_set_all test passed!")

# 测试paragraphs_set_hides_first 方法
def test_paragraphs_set_hides_first():
    document = Document()
    document.add_paragraph("标题段落")
    document.add_paragraph("正文段落")
    setter = DocumentFormatSetter(document)
    setter.paragraphs_set_hides_first(lines=[0])
    document.save("test_paragraphs_set_hides_first.docx")
    print("paragraphs_set_hides_first test passed!")

# 测试create_custom_style 和 apply_custom_style 方法
def test_custom_style_operations():
    document = Document()
    document.add_paragraph("自定义样式测试段落1")
    document.add_paragraph("自定义样式测试段落2")
    setter = DocumentFormatSetter(document)
    setter.create_custom_style("MyCustomStyle", "黑体", "Arial", 16,
                               para_alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, left_indent=8, right_indent=8,
                               space_before=8, space_after=8, line_spacing=20,
                               first_line_indent=8)
    setter.apply_custom_style(document.paragraphs[0], "MyCustomStyle")
    setter.apply_custom_style(document.paragraphs[1], "MyCustomStyle")
    document.save("test_custom_style_operations.docx")
    print("create_custom_style and apply_custom_style tests passed!")


test_pages_set()
test_sections_set()
test_para_set_indent()
test_paragraph_set()
test_paragraphs_set_all()
test_paragraphs_set_hides_first()
test_custom_style_operations()


# from document_page_number_adder import DocumentPageNumberAdder
# from docx import Document

# 测试add_page_number 方法
def test_add_page_number():
    document = Document()
    document.add_paragraph("这是一个测试文档，用于测试添加页码功能")
    adder = DocumentPageNumberAdder(document)
    adder.add_page_number()
    document.save("test_add_page_number.docx")
    print("add_page_number test passed!")