# Create:20241226
#modify codes using class and create testing code

# from docx.document import Document
from docx import Document,document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.oxml import OxmlElement
import os
import re
from pathlib import Path
import datetime


class DocumentConverter:

    def __init__(self):
        self.document = None

    def docx_old_to_docx_new(self, input_data):
        """
        将旧的docx文件内容导入，生成一个新的docx文档，防止被原文档的设置干扰。
        也能处理直接传入字符串内容的情况（当作文档内容来处理）。

        参数:
        input_data (str 或 docx.document.Document): 可以是docx文件路径字符串，也可以是已经打开的Document对象，或者是纯文本内容字符串

        返回:
        docx.document.Document: 生成的新的docx文档对象

        异常处理:
        - 如果输入的文件路径不存在，抛出FileNotFoundError异常。
        - 如果打开文件时出现其他I/O异常，向上传递相应的I/O异常。
        """
        if isinstance(input_data, str):
            if not os.path.exists(input_data):
                raise FileNotFoundError(f"指定的文件 {input_data} 不存在")
            try:
                if input_data.endswith('.docx'):
                    doc_in_text = self.get_text(input_data)
                else:
                    doc_in_text = input_data
            except IOError as e:
                raise IOError(f"读取文件 {input_data} 时出现错误: {e}")
        elif isinstance(input_data, document.Document):
            doc_in_text = '\n'.join([p.text for p in input_data.paragraphs])
        else:
            raise ValueError("输入参数类型不正确，应为 str（docx 文件路径）或 docx.document.Document 类型")

        self.document = Document()
        list_text = doc_in_text.split('\n')
        for text in list_text:
            if text.strip():
                self.document.add_paragraph(text)

        return self.document


    def get_text(self, file_name):
        """
        从docx文档中读取全部文本内容。

        参数:
        file_name (str): docx文档的文件路径

        返回:
        str: 文档中的全部文本内容，以换行符连接各段落文本

        异常处理:
        - 如果文件不存在，抛出FileNotFoundError异常。
        - 如果读取文件过程中出现其他I/O异常，向上传递相应的I/O异常。
        """
        if not os.path.exists(file_name):
            raise FileNotFoundError(f"指定的docx文件 {file_name} 不存在")
        try:
            doc = Document(file_name)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except IOError as e:
            raise IOError(f"读取docx文件 {file_name} 时出现错误: {e}")


def test_docx_old_to_docx_new_with_document_object():
    # from docx import Document
    converter = DocumentConverter()
    doc = Document()
    doc.add_paragraph("这是一段测试文本")
    new_doc = converter.docx_old_to_docx_new(doc)
    new_doc.save("test_output_docx_new_from_object.docx")
    print("docx_old_to_docx_new with document object test passed!")


test_docx_old_to_docx_new_with_document_object()
