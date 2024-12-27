#!/usr/bin/env python
# coding: utf-8



# import docx









import os
import re

import shutil

import datetime


from pathlib import Path








# import docx
# import os



def is_single_sentence(paragraph_text):
    """
    此函数用于判断一个段落是否只有一句话

    参数：
    paragraph_text (str): 要判断的段落文本

    返回：
    bool: 如果段落符合只有一句话的条件，返回 True，否则返回 False
    """
    # 去除段落两端的空白
    paragraph_text = paragraph_text.strip()
    if not paragraph_text:  # 如果段落为空，返回 False
        return False

    # 按常见的标点符号分割段落文本
    sentences = re.split(r'[。!?]', paragraph_text)
    # 如果分割后只有一个部分，并且这个部分的结尾不是逗号或者没有标点且不是因为换行导致的分割，那么认为是一句话
    if len(sentences) == 1 and not (paragraph_text.endswith(',') or (len(paragraph_text) > 0 and paragraph_text[-1] not in ['.', '!', '?', ','] and not paragraph_text.endswith('\n'))):
        return True
    else:
        return False


# In[29]:


#在python docx中实现word插入页码

#主要使用的是页脚功能。

#另外，用到了docx的oxml命令。
#https://baijiahao.baidu.com/s?id=1665454009794833226&wfr=spider&for=pc
#https://stackoverflow.com/questions/50776715/setting-pgnumtype-property-in-python-docx-is-without-effect

def AddFooterNumber(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
def InsertPageNumber(Doc):
    footer = Doc.sections[0].footer # 获取第一个节的页脚
    footer.is_linked_to_previous = True  #编号续前一节
    paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页脚居中对齐
    
    run_left = paragraph.add_run()
    run_left.text = '—  '
    font = run_left.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_left._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗
    
    run_footer=paragraph.add_run() # 添加页脚内容
    AddFooterNumber(run_footer)
    font = run_footer.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_footer._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗
    
    run_right = paragraph.add_run()
    run_right.text = '  — '
    font = run_right.font
#  font.name = 'Times New Roman'#新罗马字体
    font.name = '仿宋'#新罗马字体
    run_right._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    font.size = Pt(14)#14号字体
    font.bold = False #加粗

#EndDef


# In[30]:


def delete_paragraphs_null_lines(document):    
    '''删除docx文件的空行'''
    # 打开docx文件
    doc = document

    # 创建新文档
    new_doc = Document()

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        # 如果段落不是空行，则复制到新文档中
        if paragraph.text.strip():
            new_doc.add_paragraph(paragraph.text)

    # 保存修改后的文件
    #new_doc.save('modified_file.docx')
    return new_doc


# In[31]:




# In[32]:


# 大标题占用一行或2行
def docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1]):
    #将源文件转为新的文件，按说也可以不转，直接处置即可
    document = docx_old_to_docx_new(docx_path_in)
    #设置页面
    pages_set(document)
    #设置页边距，可以从参数文件中读取，不用修改函数，尝试使用JSON
    sections_set(document, left = 2.8, right = 2.6, top = 3.7, bottom = 3.5)
    #设置段落边距等
    para_set_indent(document)
    
    #设置正文，全部先设置为正文格式，再逐个调整各级标题
    paragraphs_set_all(document)
    #设置各级标题
    paragraphs_set_hides_first(document,lines = hides_first)
    paragraphs_set_hides_second(document)
    paragraphs_set_hides_third(document)
    #设置主送机关和落款单位、时间
    paragraphs_set_inscribe(document)
    #添加页码，设置，如果页面超过2页，则添加，另外，如果
    InsertPageNumber(document)
    
    #保存为新的文件
    document.save(docx_path_out)
    




import docx
import json



import docx

def str_to_docx(input_str):
    """
    将输入字符串转换为docx文档
    有多行的，要分行添加到docx文档的段落中

    参数：
    input_str (str): 输入字符串

    返回：
    docx.Document: 转换后的 Word 文档对象
    """
    doc = docx.Document()
    paragraphs = input_str.split('\n')
    '''当使用 input_str.split('\n') 来分割字符串时，
    如果存在连续的两个 '\n\n' ，
    那么会将这两个连续的换行符视为一个分隔符，
    分割出来的段落列表中，
    在对应位置会是一个空字符串元素。
    '''
    for paragraph in paragraphs:
        if paragraph:
            doc.add_paragraph(paragraph)
    return doc


# In[34]:


# import os
# import shutil

def move_file_to_new_directory(absolute_file_path, new_directory_name):
    '''python,编写一个函数或类，
    参数为一个文件的绝对目录和新目录的名字（当前目录下的新目录），
    查看有无新目录，没有则新建，有则忽略。
    然后将这个文件移动到这个新目录中。
    '''
    current_directory = os.path.dirname(absolute_file_path)
    new_directory_path = os.path.join(current_directory, new_directory_name)
    if not os.path.exists(new_directory_path):
        os.mkdir(new_directory_path)
    shutil.move(absolute_file_path, new_directory_path)
    
    


def copy_file_to_new_directory(absolute_file_path, new_directory_name):
    '''
    python,编写一个函数或类，
    参数为一个文件的绝对目录和新目录的名字（当前目录下的新目录），
    查看有无新目录，没有则新建，有则忽略。然后将这个文件复制到这个新目录中。
    '''
    current_directory = os.path.dirname(absolute_file_path)
    new_directory_path = os.path.join(current_directory, new_directory_name)
    if not os.path.exists(new_directory_path):
        os.mkdir(new_directory_path)
    shutil.copy(absolute_file_path, new_directory_path)


# In[35]:


# import os

def create_directory_if_not_exists(directory_path, folder_name):
    """
    此函数用于检查指定目录中是否存在给定的文件夹，如果不存在则创建

    参数：
    directory_path (str): 要检查的目录路径
    folder_name (str): 要检查和创建的文件夹名称

    返回：
    None
    """
    target_directory = os.path.join(directory_path, folder_name)
    if not os.path.exists(target_directory):
        os.makedirs(target_directory)
    return target_directory


# In[36]:


#docx_to_offical_format(docx_path_in, docx_path_out)


# In[37]:


#批量修改为公文格式，并保存在独立的文件夹中，但复用不方便
def modify_offical_docxs(folder_path_in):
    '''
    将指定文件夹中的txt和docx文件，
    转为标准公文格式（带主送和落款，标题占用2行）
    新文件保存在folder_path_in的兄弟目录“新生成的标准公文格式文件”中
    '''
    date = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    txt_to_docx_all(folder_path_in)
    #docxs = []
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.docx'):
            docx_path_in = os.path.join(folder_path_in, file_name_in)
            #docxs.append(docx_path_in)
#             folder_path = os.path.dirname(folder_path_in) #将folder_path变为局部变量，防止出错,只生成了当前目录，不明所以
            folder_path = str(Path(folder_path_in).parent).replace('\\','/')
            folder_path_out = create_directory_if_not_exists(folder_path,   '生成的标准公文格式文件/')
            docx_path_out = os.path.join(folder_path_out , (date + file_name_in))  #此处有问题，使用了全局变量folder_path，已修改
        
            docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])
            os.startfile(docx_path_out)
            
            
#批量修改为公文格式，直接在原目录生成新的同名文件，增加标准公文格式字样，或其他字样
def modify_offical_docxs_in_current_directory(folder_path_in, folder_relative_path_out = '生成的文件',new_directory_name = '旧',open_new = True):
    # folder_relative_path_out = None 是否可以？
    '''
    将指定文件夹中的txt和docx文件，
    转为标准公文格式（带主送和落款，标题占用2行）
    新文件保存在folder_path_in的子目录folder_relative_path_out中，没有则生成一个
    将原文件保存在子目录new_directory_name中
    '''
    date = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    folder_path_out = create_directory_if_not_exists(folder_path_in,  folder_relative_path_out)

    #不要生成中间文件夹了吧？
    '''
    读取folder_path_in中的txt和docx文件
    将txt转为docx
    将docx们生成带主送和落款的标准公文格式docx
    '''
    # create_directory_if_not_exists(folder_path_in,  new_directory_name)

    txt_to_docx_all(folder_path_in)  #这个要看看保存到了哪里，是同目录下吧？
    #docxs = []
    for file_name_in in os.listdir(folder_path_in):
        if file_name_in.endswith('.docx'):
            docx_path_in = os.path.join(folder_path_in, file_name_in)
            #docxs.append(docx_path_in)
            docx_path_out = os.path.join(folder_path_out , (date + file_name_in))
            # copy_file_to_new_directory(docx_path_in, new_directory_name)
        
            docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])
#             docx_to_offical_format(docx_path_in, docx_path_in,hides_first =  [0,1])  #直接覆盖原文件，可否？
#             os.startfile(docx_path_in) 不需要打开此文件

            print(f'Format {docx_path_in} to {docx_path_out}\n')
            if open_new:
                os.startfile(docx_path_out)


# In[38]:


# import docx

import re

def is_single_sentence(paragraph_text):
    """
    此函数用于判断一个段落是否只有一句话

    参数：
    paragraph_text (str): 要判断的段落文本

    返回：
    bool: 如果段落符合只有一句话的条件，返回 True，否则返回 False
    """
    paragraph_text = paragraph_text.strip()
    if not paragraph_text:
        return False

    # 去除换行符
    paragraph_text = paragraph_text.replace('\n', '')

    # 按句号、感叹号、问号分割
    sentences = re.split(r'[。!?]', paragraph_text)
    sentences = [item for item in sentences if item!= '']  #要删除列表中为空''的部分
    print(sentences)
    # 检查分割后的部分数量以及最后一个部分是否只是标点后的剩余部分（如：“这是一句话。” 分割后为 ["这是一句话", ""] ）
    if len(sentences) == 1 and not sentences[0] == '':
        return True
    else:
        return False


# In[ ]:





# In[39]:


# import os

def find_docx_files(directory):
    '''
    python编写一个函数或类，参数为一个目录，
    查找一个目录中所有以.docx为后缀的文件，
    将全部文件的绝对路径保存到一个列表中，
    并返回。'''
    docx_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                docx_files.append(file_path)
    return docx_files

import os  # 导入 os 模块，用于文件和目录操作

def find_files_by_suffix(directory, suffix, include_subdirectories=False):  # 定义函数 find_files，接受目录、后缀和是否包含子目录的参数
    """
    此函数用于根据给定的目录和后缀（或后缀列表）查找文件，并根据是否包含子目录的参数决定是否遍历子目录，将符合条件的文件的绝对路径保存到列表中返回

    参数：
    directory (str): 要搜索的目录路径
    suffix (str 或 list): 文件的后缀或后缀列表,或后缀元组
    include_subdirectories (bool, 可选): 是否包含子目录，默认为 False

    返回：
    list: 包含符合条件文件的绝对路径的列表
    
    
    """
    file_paths = []  # 创建一个空列表用于存储文件路径
    for root, dirs, files in os.walk(directory):  # 使用 os.walk 遍历目录及其子目录
        """
        root: 当前遍历的目录路径
        dirs: 当前目录下的子目录列表
        files: 当前目录下的文件列表
        """
        if not include_subdirectories and root!= directory:  # 如果不包含子目录且当前目录不是指定目录，则跳过
            continue
        for file in files:  # 遍历当前目录下的文件
            if file.endswith(tuple(suffix)) or (isinstance(suffix, list) and any(file.endswith(suffix_item) for suffix_item in suffix)):  # 判断文件后缀是否符合条件
                """
                当 suffix 是列表时，将其转换为元组传递给 endswith 方法
                
                这里进行后缀的匹配判断。
                `file.endswith(suffix)` 检查文件的后缀是否与给定的单个后缀 `suffix` 相匹配。
                如果 `suffix` 是一个列表，`isinstance(suffix, list) and any(file.endswith(suffix_item) for suffix_item in suffix)` 会检查文件的后缀是否与列表中的任何一个后缀相匹配。
                `any()` 函数用于判断是否存在至少一个满足条件的情况。
                """
                file_paths.append(os.path.abspath(os.path.join(root, file)))  # 将符合条件的文件的绝对路径添加到列表中
                """
                如果文件的后缀符合条件，使用 `os.path.abspath(os.path.join(root, file))` 计算文件的绝对路径，并将其添加到 `file_paths` 列表中。
                `os.path.join(root, file)` 用于组合目录路径 `root` 和文件名 `file` ，形成完整的文件路径。
                `os.path.abspath()` 用于获取该路径的绝对路径。
                """
    return file_paths  # 返回包含文件绝对路径的列表


# In[40]:


import docx
import os
import win32com.client as win32
#可以获取页面，但是依赖于安装了word,慎用
def get_word_page_count(document_path):
    document = docx.Document(document_path)
    abs_path = os.path.abspath(document_path)

    # 保存文档为临时文件
    temp_path = os.path.join(os.path.dirname(abs_path), "~temp.docx")
    document.save(temp_path)

    # 打开临时文件并获取页数
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(temp_path)
    pages = doc.BuiltInDocumentProperties(14).Value
    doc.Close()
    word.Quit()

    # 删除临时文件
    os.remove(temp_path)

    return pages

# document_path = docx_path_out_txt
# page_count = get_word_page_count(document_path)
# print("Word文档的页数：", page_count)
# 没想到获取个word页数这么简单的需求，实现起来还挺难。

# 1.根据网上说法：word文档中不保存总页数，只是在渲染时才根据文字和表格等去摆放到一页上，摆放不下时才展现下一页。所以从docx中取不到总页数

# 2.因此，仅用docx库不行，还需要用win32com。而win32com实际是：pip install pypiwin32   并且一定记得要重启python（重启仍不行可以试试https://blog.csdn.net/weixin_43149311/article/details/120806116  ，我重启IDLE就行了）

# 要获取Word文档的页数，我们可以使用以下步骤：

# 首先，将Word文档保存为临时文件。
# 然后，使用win32com.client模块打开该临时文件。
# 接下来，获取打开的Word应用程序的ActiveDocument。
# 最后，通过BuiltInDocumentProperties属性获取文档的页数。

# 补充：今天执行上述代码出错： 'win32com.gen_py....' has no attribute 'CLSIDToClassMap'

#  解决办法：删除目录<code>C:\Users\你的用户名\AppData\Local\Temp\gen_py\3.X中的缓存文件夹00020905-0000-0000-C000-000000000046x0x8x5即可，重新执行上述代码便不再报错。

 

# 参考：https://blog.51cto.com/u_16175516/6850938

# https://blog.csdn.net/u011840075/article/details/124696371

# https://blog.csdn.net/weixin_43149311/article/details/120806116

# https://blog.51cto.com/mouday/5018378

# https://blog.csdn.net/ericatardicaca/article/details/90721909


# In[ ]:





# In[ ]:





# In[41]:


#获取当前时间
date = datetime.datetime.now().strftime("%Y%m%d")
#当前主文件夹
# folder_path = 'D:/工作自动化/调整公文格式带主送落款'#最后面还有个/，别遗漏
# folder_path = 'D:/工作自动化/生成公文格式20240730'#最后面还有个/，别遗漏

folder_path = Path.cwd()
folder_path = os.path.normpath(os.getcwd()).replace('\\', '/') + '/'
#也可以获取当前文件夹

folder_path_in = folder_path + '/未调整的源文件/'
# /folder_path_in = folder_path / Path( '未调整的源文件')
#文件夹中只保存1个docx文件
# docxs = []
# for file_name_in in os.listdir(folder_path_in):
#     if file_name_in.endswith('.docx'):
#         docx_path_in = os.path.join(folder_path_in, file_name_in)
#         docxs.append(docx_path_in)
#         docx_path_out = os.path.join(folder_path , ( '标准公文格式文件/'+ date + file_name_in))

# print(docx_path_out)


# In[42]:


# 测试用，删除即可
docx_path_out_str = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111str生成的.docx'
docx_path_out_txt = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111txt生成的.docx'
docx_path_out_txt222 = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\222txt生成的.docx'
docx_path_out_docx = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111doc生成的.docx'
# doc = str_to_docx(input_file)
txt_path_in_tmp = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111.txt'
txt_path_in_tmp222 = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\222.txt'
docx_path_in_tmp = 'D:\\工作自动化\\生成公文格式20240730\\未调整的源文件\\111.docx'



# In[43]:


# import pprint
# pprint.pprint(find_files_by_suffix(folder_path_in, 'docx', include_subdirectories=False))


# In[44]:


#将一个docx\txt\srt作为输入，然后生成公文格式文件，单独调用这个最方便
# file_to_docx(path_in, docx_path_out, hides_first=[0, 1])


# In[45]:


#    将指定文件夹中的txt和docx文件，
#     转为标准公文格式（带主送和落款，标题占用2行）
#     新文件保存在folder_path_in的兄弟目录“新生成的标准公文格式文件”中

# modify_offical_docxs(folder_path_in)


# path_tmp = 'D:/code_jupyter/文本转公文格式全自动20241024/kwkw哈哈'
# modify_offical_docxs(path_tmp)

# In[46]:


#os.startfile(docx_path_out)


# In[47]:


# modify_offical_docxs_in_current_directory(folder_path_in, folder_relative_path_out = '生成的文件',new_directory_name = '旧')


# In[48]:


# modify_offical_docxs(folder_path_in)


# In[49]:


#把一个文件转换，调用这个的会比较多，但是不能直接把字符串转换，这样多少会降低效率
# docx_to_offical_format(docx_path_in, docx_path_out,hides_first =  [0,1])


# In[51]:


#可调用的函数为：
# 1.标准公文，一行标题，主送、落款单位和日期# 
#生成标准公文格式的，带主送和落款，可选首行的行数（1行或2行），后面的可能有bug,主送那里
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0, 1])  

# 2.标准公文，2行标题，主送、落款单位和日期
# 3.普通公文，1行标题， 无主送、落款单位和日期
# file_to_docx(txt_path_in_tmp, docx_path_out_txt, hides_first=[0],have_inscribe = False)
#4.条例，个别需要修改的地方，如加粗第几章，第几条，居中的地方
#5.缩小版，打印较多的材料用
#6.领导发言（第二行居中楷体的那种,也可能有日期，即标题下有2行居中楷体）

# file_to_docx(txt_path_in_tmp222, docx_path_out_txt222, hides_first=[0],have_inscribe = False,date_and_name_under_title_flag = True,date_and_name_under_title_numbers = [1])


# In[ ]:


# file_to_docx(
#     txt_path_in_tmp222, 
#     docx_path_out_txt222,
#     hides_first=[0],
#     have_inscribe = False,
#     date_and_name_under_title_flag = True,
#     date_and_name_under_title_numbers = [1]
# )









